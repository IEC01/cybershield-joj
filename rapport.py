from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

doc = Document()

# ── Marges ──────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Fonctions utilitaires ────────────────────────────────
def titre_principal(texte):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p

def sous_titre(texte):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    return p

def titre_section(texte):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.space_after = Pt(6)
    return p

def paragraphe(texte):
    p = doc.add_paragraph(texte)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def encadre_info(texte, couleur_fond=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(texte)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p

def separateur():
    p = doc.add_paragraph("─" * 72)
    for run in p.runs:
        run.font.size  = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def placeholder_capture(legende):
    """Ajoute un espace réservé pour une capture d'écran"""
    p = doc.add_paragraph()
    run = p.add_run(f"[ CAPTURE D'ÉCRAN — {legende} ]")
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.italic    = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cadre visuel
    p2 = doc.add_paragraph()
    run2 = p2.add_run("┌" + "─" * 60 + "┐\n│" + " " * 60 + "│\n│" +
                       "  Insérer capture ici (Insertion > Images)  ".center(60) +
                       "│\n│" + " " * 60 + "│\n└" + "─" * 60 + "┘")
    run2.font.size  = Pt(9)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run2.font.name  = "Courier New"
    p2.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

titre_principal("🛡 CYBERSHIELD JOJ")
sous_titre("Plateforme de Détection d'Attaques — Jeux Olympiques de la Jeunesse")
doc.add_paragraph()
sous_titre("Dakar 2026")
doc.add_paragraph()
separateur()
doc.add_paragraph()

# Infos document
table_garde = doc.add_table(rows=4, cols=2)
table_garde.alignment = WD_TABLE_ALIGNMENT.CENTER
infos = [
    ("Projet",   "CyberShield JOJ — Hackathon Dakar 2026"),
    ("Date",     datetime.now().strftime("%d %B %Y")),
    ("Version",  "1.0 — Prototype Hackathon"),
    ("Équipe",   "À compléter"),
]
for i, (cle, val) in enumerate(infos):
    row = table_garde.rows[i]
    row.cells[0].text = cle
    row.cells[1].text = val
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. RÉSUMÉ EXÉCUTIF
# ══════════════════════════════════════════════════════════
titre_section("1. Résumé exécutif")
separateur()
paragraphe(
    "CyberShield JOJ est une plateforme de cybersécurité développée dans le cadre "
    "du hackathon des Jeux Olympiques de la Jeunesse Dakar 2026. Elle assure la "
    "surveillance en temps réel des systèmes critiques et détecte automatiquement "
    "les cyberattaques grâce à un moteur d'intelligence artificielle."
)
paragraphe(
    "Les Jeux Olympiques de la Jeunesse représentent une cible de choix pour les "
    "cyberattaquants : billetterie en ligne, accréditations des athlètes, diffusion "
    "des résultats en direct. Une attaque réussie pendant l'événement causerait un "
    "préjudice d'image mondial. CyberShield JOJ répond à ce risque."
)

doc.add_paragraph()
titre_section("2. Problème adressé")
separateur()
paragraphe("Les systèmes critiques des JOJ Dakar 2026 exposés aux cybermenaces :")

menaces = [
    ("Billetterie",      "billetterie.joj2026.sn",    "DDoS, fraude"),
    ("Accréditations",   "accreditations.joj2026.sn", "Brute Force, usurpation"),
    ("Résultats sportifs","resultats.joj2026.sn",     "Falsification, injection SQL"),
    ("Site officiel",    "site-officiel.joj2026.sn",  "Defacement, DDoS"),
]

table_menaces = doc.add_table(rows=5, cols=3)
table_menaces.style = "Table Grid"
headers = ["Système", "Domaine", "Menaces principales"]
for j, h in enumerate(headers):
    cell = table_menaces.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    from docx.oxml.ns import qn
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1E40AF")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

for i, (sys, dom, threat) in enumerate(menaces):
    row = table_menaces.rows[i + 1]
    row.cells[0].text = sys
    row.cells[1].text = dom
    row.cells[2].text = threat
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  3. ARCHITECTURE
# ══════════════════════════════════════════════════════════
titre_section("3. Architecture technique")
separateur()
paragraphe(
    "CyberShield JOJ repose sur une architecture en pipeline à 5 couches, "
    "du trafic réseau brut jusqu'au tableau de bord de sécurité :"
)

pipeline = [
    ("Couche 1", "Trafic Web",       "Requêtes HTTP/HTTPS entrantes sur les 4 systèmes JOJ"),
    ("Couche 2", "IDS / Snort",      "Détection par règles — seuils de trafic, signatures d'attaques"),
    ("Couche 3", "Collecteur",       "Extraction des features réseau : req_rate, bytes, port, durée"),
    ("Couche 4", "Moteur IA",        "Isolation Forest — score d'anomalie entre 0 et 1"),
    ("Couche 5", "Dashboard SOC",    "Alertes temps réel, métriques, simulateur d'attaques"),
]

table_arch = doc.add_table(rows=6, cols=3)
table_arch.style = "Table Grid"
for j, h in enumerate(["Couche", "Composant", "Rôle"]):
    cell = table_arch.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1E40AF")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

for i, (c, comp, role) in enumerate(pipeline):
    row = table_arch.rows[i + 1]
    row.cells[0].text = c
    row.cells[1].text = comp
    row.cells[2].text = role
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
titre_section("Stack technologique")
tech = [
    "Backend    : Python 3 + FastAPI + WebSocket",
    "IA         : Scikit-learn — Isolation Forest + StandardScaler",
    "Frontend   : HTML5 / CSS3 / JavaScript vanilla",
    "Protocole  : WebSocket temps réel (push serveur → client)",
    "Déploiement: Uvicorn ASGI sur Linux / Kali",
]
for t in tech:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(t)
    run.font.size = Pt(11)
    run.font.name = "Courier New"

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  4. CAPTURES D'ÉCRAN
# ══════════════════════════════════════════════════════════
titre_section("4. Captures d'écran du prototype")
separateur()
paragraphe(
    "Les captures ci-dessous illustrent le fonctionnement de CyberShield JOJ "
    "en conditions de démonstration. Pour insérer vos captures : "
    "cliquez sur un encadré → Insertion → Images → sélectionnez votre fichier."
)
doc.add_paragraph()

captures = [
    "Dashboard principal — trafic normal, score IA < 0.50",
    "Simulation attaque DDoS — score IA 0.55+, bannière rouge",
    "Panneau Alertes récentes — liste des anomalies détectées",
    "Console F12 — WebSocket connecté, logs JSON en temps réel",
    "Terminal — serveur uvicorn + modèle IA entraîné",
]
for cap in captures:
    encadre_info(f"Capture : {cap}")
    placeholder_capture(cap)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. MODULE IA
# ══════════════════════════════════════════════════════════
titre_section("5. Module d'intelligence artificielle")
separateur()
paragraphe(
    "Le moteur IA de CyberShield JOJ utilise l'algorithme Isolation Forest, "
    "un modèle d'apprentissage non supervisé particulièrement adapté à la "
    "détection d'anomalies dans les flux réseau."
)

titre_section("Principe de fonctionnement")
principes = [
    "Phase d'entraînement : le modèle apprend le comportement normal du réseau "
    "à partir de 1 000 échantillons de trafic légitime simulé.",
    "Phase d'analyse : chaque requête est projetée dans un espace à 4 dimensions "
    "(req_rate, bytes, port, durée) puis comparée au modèle.",
    "Score d'anomalie : une valeur entre 0 et 1 est calculée via une fonction "
    "sigmoïde. Un score > 0.50 déclenche une alerte.",
    "Avantage clé : contrairement aux systèmes à règles statiques, l'IA s'adapte "
    "au profil de trafic de chaque environnement.",
]
for p_text in principes:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(p_text)
    run.font.size = Pt(11)

doc.add_paragraph()
titre_section("Résultats observés lors des tests")
resultats = doc.add_table(rows=4, cols=3)
resultats.style = "Table Grid"
for j, h in enumerate(["Type de trafic", "Score IA", "Verdict"]):
    cell = resultats.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1E40AF")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

data_resultats = [
    ("Trafic normal",     "0.449", "Normal — aucune alerte"),
    ("Attaque DDoS",      "0.552", "Anomalie détectée — IP bloquée"),
    ("Port Scan",         "0.546", "Anomalie détectée — IP surveillée"),
]
for i, (t, s, v) in enumerate(data_resultats):
    row = resultats.rows[i + 1]
    row.cells[0].text = t
    row.cells[1].text = s
    row.cells[2].text = v
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  6. FONCTIONNALITÉS
# ══════════════════════════════════════════════════════════
titre_section("6. Fonctionnalités du prototype")
separateur()

fonctionnalites = [
    ("Surveillance temps réel",
     "Flux WebSocket continu entre le serveur et le dashboard. "
     "Chaque requête réseau est analysée et affichée en moins d'une seconde."),
    ("Détection IA des anomalies",
     "Isolation Forest entraîné sur du trafic normal. Score entre 0 et 1 "
     "affiché en temps réel avec barre de progression colorée."),
    ("Simulateur d'attaques",
     "4 types d'attaques simulables depuis le dashboard : DDoS, Brute Force, "
     "Port Scan, SQL Injection. Retour au trafic normal automatique après 15s."),
    ("Alertes visuelles et sonores",
     "Bannière rouge + notification sonore (double bip) à chaque anomalie. "
     "Panneau des alertes récentes avec horodatage et score IA."),
    ("Métriques en direct",
     "4 indicateurs clés : requêtes/min, anomalies détectées, IPs bloquées, "
     "score de sécurité global calculé en temps réel."),
    ("Contexte JOJ Dakar 2026",
     "4 systèmes critiques modélisés : billetterie, accréditations, résultats "
     "sportifs, site officiel — domaines .joj2026.sn."),
]

for titre, desc in fonctionnalites:
    p = doc.add_paragraph()
    run_titre = p.add_run(f"{titre} : ")
    run_titre.bold = True
    run_titre.font.size = Pt(11)
    run_titre.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  7. GUIDE DE DÉMO
# ══════════════════════════════════════════════════════════
titre_section("7. Guide de démonstration — 5 minutes")
separateur()
paragraphe(
    "Script de présentation devant le jury. Chaque étape est chronométrée "
    "pour tenir dans le temps imparti."
)
doc.add_paragraph()

etapes = [
    ("30 sec", "Lancement",
     "Lancer uvicorn main:app --reload dans le terminal. "
     "Ouvrir http://localhost:8000 dans le navigateur. "
     "Le dashboard se connecte automatiquement — point vert visible."),
    ("60 sec", "Trafic normal",
     "Montrer les logs qui défilent avec des scores < 0.50. "
     "Expliquer les 4 systèmes JOJ surveillés. "
     "Score de sécurité à 100%."),
    ("90 sec", "Simulation DDoS",
     "Cliquer sur le bouton DDoS dans le simulateur. "
     "Montrer le score IA monter à 0.55+. "
     "Bannière rouge + bip sonore. Anomalies comptabilisées."),
    ("60 sec", "Expliquer l'IA",
     "Isolation Forest apprend le comportement normal. "
     "Score sigmoïde entre 0 et 1 — pas un simple binaire. "
     "S'adapte à l'environnement, contrairement aux règles statiques."),
    ("40 sec", "Conclusion JOJ",
     "CyberShield JOJ protège les 4 systèmes critiques de Dakar 2026. "
     "Open source, déployable en 24h. "
     "Une cyberattaque pendant les JOJ = désastre d'image mondiale."),
]

table_demo = doc.add_table(rows=6, cols=3)
table_demo.style = "Table Grid"
for j, h in enumerate(["Durée", "Étape", "Ce que vous dites / faites"]):
    cell = table_demo.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1E40AF")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

for i, (dur, step, desc) in enumerate(etapes):
    row = table_demo.rows[i + 1]
    row.cells[0].text = dur
    row.cells[1].text = step
    row.cells[2].text = desc
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  8. CONCLUSION
# ══════════════════════════════════════════════════════════
titre_section("8. Conclusion et perspectives")
separateur()
paragraphe(
    "CyberShield JOJ démontre qu'une solution de cybersécurité intelligente "
    "et temps réel peut être construite rapidement avec des outils open source. "
    "Le prototype présenté lors du hackathon couvre l'ensemble de la chaîne : "
    "collecte du trafic, analyse IA, visualisation et réponse aux incidents."
)
paragraphe(
    "Pour une version production, les évolutions prioritaires seraient : "
    "intégration avec un vrai IDS (Snort/Zeek), base de données pour "
    "l'historique des alertes, authentification du dashboard, et déploiement "
    "sur un serveur dédié avec accès depuis mobile."
)

doc.add_paragraph()
separateur()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_fin = p_fin.add_run("CyberShield JOJ — Hackathon Dakar 2026")
run_fin.font.size  = Pt(11)
run_fin.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run_fin.font.italic    = True

# ── Sauvegarde ───────────────────────────────────────────
nom_fichier = f"CyberShield_JOJ_Rapport_{datetime.now().strftime('%Y%m%d')}.docx"
chemin      = os.path.join(os.path.expanduser("~"), nom_fichier)
doc.save(chemin)
print(f"Rapport généré : {chemin}")
